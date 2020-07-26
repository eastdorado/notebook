#!/usr/bin/python3
# -*- coding: utf-8 -*-

# @项目名称: python
# @File    : word2pdf.py
# @Time    : 2020/2/20 15:20
# @Author  : big
# @Email   : shdorado@126.com

import sys
import os
import timeit
import multiprocessing
import json
from PyQt5 import QtCore, QtGui, QtWidgets
from utilities import Utils, AnimWin
from ui_word2pdf import Ui_MainWindow
import docx
import tkinter
import tkinter.filedialog


class MultiPro:
    def __init__(self):
        cpu_num = Utils.GetCpuInfo()[0]
        # print(cpu_num)
        # self.pool = multiprocessing.Pool()
        self.pool = multiprocessing.Pool(processes=cpu_num)  # 根据cpu个数创建进程池里的进程个数
        # 如果不设置参数,函数会跟根据计算机的实际情况来决定要运行多少个进程，自己设置要考虑计算机的性能

    def run(self, fun_pro, args, fun_back=None):
        # print('run', args)
        result = self.pool.apply_async(fun_pro, args, fun_back)
        # print(results.get())
        # time.sleep(3)
        return result

    def end(self):
        self.pool.close()
        self.pool.join()


class MainWindow(QtWidgets.QMainWindow, Ui_MainWindow):
    def __init__(self):
        super(MainWindow, self).__init__()
        self.data_path = {}
        # self.dir_src = None
        # self.dir_dst = None

        self.count = 0  # 转换失败的文件数
        self.checked = False

        self.files_src = []
        self.files_dst = []
        self.files_all = []

        try:
            with open("./res/word2pdf.json", 'r') as load_f:
                self.data_path = json.load(load_f)
        except IOError as e:
            self.data_path = {'dir_src': '', 'dir_dst': ''}
            print(e)
        finally:
            # print(self.data_path)
            pass

        self.setupUi(self)
        # TODO 修改原始控件
        self.resize(1000, 600)

        self.setStyleSheet('font-size:18px;font-weight:bold;font-family:Roman times;')

        # qss = Utils.readQss(r'E:\python\res\style\list_table.qss')
        qss = 'QListWidget{background-color:transparent; color:black; border:1px solid gray;}' \
              'QListWidget::Item{padding-top:-2px; padding-bottom:-1px;}' \
              'QListWidget::Item:hover{background:skyblue;padding-top:0px; padding-bottom:0px;}' \
              'QListWidget::item:selected{background:lightgreen; color:red;}' \
              'QListWidget::item:selected:!active{border-width:0px;background:lightgray;}'
        self.listWidget_src.setStyleSheet(qss)
        # self.listWidget_src.setStyleSheet("background-color:transparent; color:red;font-size:22px;")

        self.listWidget_dst.setStyleSheet("background-color:transparent")
        self.listWidget_src.setSelectionMode(QtWidgets.QAbstractItemView.MultiSelection)  # 点击多选
        self.label_src.setStyleSheet('color:rgb(10,10,210)')
        self.label_dst.setStyleSheet('color:rgb(10,10,210)')
        self.checkBox.setStyleSheet('font-size:16px;color:rgb(210,10,210)')

        self.statusbar.setStyleSheet('color:rgb(210,10,10)')
        self.statusbar.showMessage('总文件:0    选中:0', 0)  # 状态栏本身显示的信息 第二个参数是信息停留的时间，单位是毫秒，默认是0（0表示在下一个操作来临前一直显示）

        self.successNum = QtWidgets.QLabel('转换成功:0    ')
        self.failNum = QtWidgets.QLabel('失败:0    ')
        self.statusbar.addPermanentWidget(self.successNum, stretch=0)
        self.statusbar.addPermanentWidget(self.failNum, stretch=0)

        self.flushDir()

    def fun_callback(self, ret_file_dst_name):
        '''
        进程的回调函数
        :param ret_file_dst_name: 返回的文件名
        :return:
        '''

        if ret_file_dst_name:
            self.files_dst.append(ret_file_dst_name)
            item = QtWidgets.QListWidgetItem(ret_file_dst_name)
            self.listWidget_dst.addItem(item)
            self.successNum.setText(f'转换成功:{self.listWidget_dst.count()}    ')
        else:
            self.count += 1  # 失败次数
            self.failNum.setText(f'失败:{self.count}    ')

    @staticmethod
    def fun_process(file_src_name, dir_src, dir_dst):
        # print(dir_dst)
        # print('son_process')
        # print(file_src_name)
        file_dst_name = None

        # file_src = os.path.join(dir_src, file_src_name)  # 构造完整路径
        file_src = '%s/%s'%(dir_src, file_src_name)
        portion = os.path.splitext(file_src_name)  # 分离文件名与扩展名

        # 如果文件是 word文档
        if portion[1] == '.doc' or portion[1] == '.docx':
            # 重新组合文件名和后缀名
            file_dst_name = portion[0] + '.pdf'
            # file_dst = os.path.join(dir_dst, file_dst_name)  # 构造完整路径
            file_dst = '%s/%s'%(dir_dst, file_dst_name)
            print(file_src, file_dst)
            if Utils.word2Pdf(file_src, file_dst) != 0:
                file_dst_name = None
        print(file_dst_name)
        return file_dst_name

    def flushDir(self):
        dir_path = self.data_path['dir_src']
        if not dir_path:
            AnimWin('没有源目录')
            return

        tmp = Utils.getSubStr(dir_path, 10)
        # print(tmp)
        self.label_src.setText(u'源文件目录:' + tmp)

        self.files_all.clear()
        self.listWidget_src.clear()
        self.checked = False

        self.files_all = Utils.files_in_dir(dir_path, ['.doc', '.docx'])
        if not self.files_all:
            return
        Utils.sort_nicely(self.files_all)

        for each in self.files_all:
            item = QtWidgets.QListWidgetItem(each)
            self.listWidget_src.addItem(item)

        self.statusbar.showMessage(f'总文件:{len(self.files_all)}    选中:0', 0)

    def slot_tools(self, action):
        try:
            # action = QtWidgets.QAction()
            text = action.text()
            # print(text)
            if text == '刷新':
                self.flushDir()

            if text == '打开':
                # fname = tkinter.filedialog.askdirectory()
                # print(fname)
                # return
                # dialog = QtWidgets.QFileDialog(self)
                # # dialog.setAcceptMode(QtWidgets.QFileDialog.AcceptOpen)
                # # dialog.setDirectory(self.data_path['dir_src'])
                # dialog.setFileMode(QtWidgets.QFileDialog.Directory)
                # # dialog.setOption(QtWidgets.QFileDialog.ShowDirsOnly, True)
                # dialog.setViewMode(QtWidgets.QFileDialog.List)
                #
                # dir_path = dialog.getExistingDirectory()

                dir_path = QtWidgets.QFileDialog.getExistingDirectory(
                    self, "选取源文件夹", self.data_path['dir_src'],
                    QtWidgets.QFileDialog.ShowDirsOnly)  # 起始路径
                if dir_path:
                    self.data_path['dir_src'] = dir_path
                    self.flushDir()

            elif text == '保存':
                dir_path = QtWidgets.QFileDialog.getExistingDirectory(
                    self, "选取保存目录", self.data_path['dir_dst'])  # 起始路径
                if dir_path:
                    self.data_path['dir_dst'] = dir_path
                    tmp = Utils.getSubStr(dir_path, 10)
                    self.label_dst.setText(u'源文件目录:' + tmp)

            elif text == 'word 文件合并':
                '''合并word文档'''
                self.files_src.clear()
                self.files_dst.clear()
                self.listWidget_dst.clear()

                self.files_src = set([item.text() for item in self.listWidget_src.selectedItems()])
                if len(self.files_src) <= 0:
                    AnimWin('  请选择需要转换的文件！', self)
                    return

                path_all = []
                for each in self.files_src:
                    # print(each)
                    ps = os.path.join(self.data_path['dir_src'], each)  # 构造完整路径
                    path_all.append(ps)
                pd = os.path.join(self.data_path['dir_dst'], r'merged.doc')  # 构造完整路径
                # print(path_all)
                Utils.mergewords(path_all, pd)
                self.fun_callback(r'merged.doc')

            elif text == 'word 转 pdf':
                if not self.data_path['dir_src']:
                    AnimWin(' 请选择源文件所在目录！', self)
                    return
                if not self.data_path['dir_dst']:
                    AnimWin(' 请选择保存文件的目录！', self)
                    return
                if not self.files_all:
                    AnimWin(' 未发现源文件！', self)
                    return

                self.files_src.clear()
                self.files_dst.clear()
                self.listWidget_dst.clear()

                self.files_src = set([item.text() for item in self.listWidget_src.selectedItems()])
                if len(self.files_src) <= 0:
                    AnimWin('  请选择需要转换的文件！', self)
                    return

                # print("-----start-----")
                # start = timeit.default_timer()
                # cpu_num = Utils.GetCpuInfo()[0]
                # pool = multiprocessing.Pool(cpu_num)
                # for each in self.files_src:
                #     # results = pool.apply_async(self.test, (i,))
                #     args = (each, self.data_path['dir_src'], self.data_path['dir_dst'])
                #     pool.apply_async(self.fun_process, args, callback=self.fun_callback)
                # pool.close()
                # pool.join()
                # elapsed = (timeit.default_timer() - start)
                # print("-----end-----")
                # print(f'用时:{elapsed}秒')

                print("-----start-----")
                start = timeit.default_timer()

                for each in self.files_src:
                    # print(each)
                    ret = self.fun_process(each, self.data_path['dir_src'], self.data_path['dir_dst'])
                    self.fun_callback(ret)
                elapsed = (timeit.default_timer() - start)
                print("-----end-----")
                print(f'用时:{elapsed}秒')

        except Exception as e:
            print(e)

    def slot_selected(self):
        # print(len(self.files_all))
        self.statusbar.showMessage(f'总文件:{len(self.files_all)}    '
                                   f'选中:{len(self.listWidget_src.selectedItems())}', 0)

    def slot_checked(self, checked):
        self.checked = checked
        for row in range(self.listWidget_src.count()):
            item = self.listWidget_src.item(row)
            item.setSelected(self.checked)
            # if self.checked:
            #     item.setTextColor(QtGui.QColor('blue'))
            # else:
            #     item.setForeground(QtGui.QColor('black'))
            # item.setBackgroundColor(QtGui.QColor('white'))

    def closeEvent(self, event):
        # 清理一些 自己需要关闭的东西
        try:
            with open('./res/word2pdf.json', 'w') as f:
                json.dump(self.data_path, f)
            event.accept()  # 界面的关闭   但是会有一些时候退出不完全    需要调用 os 的_exit 完全退出
        except Exception as e:
            print(e)
        finally:
            os._exit(5)

    def resizeEvent(self, event):
        palette = QtGui.QPalette()
        pix = QtGui.QPixmap('res/images/background.jpg')
        pix = pix.scaled(self.width(), self.height())
        palette.setBrush(QtGui.QPalette.Background, QtGui.QBrush(pix))
        self.setPalette(palette)


if __name__ == '__main__':
    # app = QtWidgets.QApplication(sys.argv)
    # # win = MainWindow()
    # win.show()
    # sys.exit(app.exec_())
    ll = Utils.getFiles(r'C:\Users\big\Desktop\弘光2020年暑期数学作业', ['.docx'])
    print(len(ll), ll)
    for i in range(len(ll)):
        file = ll[i]
        filep, dd = os.path.splitext(file)
        Utils.word2Pdf(file, f'{filep}.pdf')
    # Utils.word2Pdf(r'D:/下载/数理化讲义/第3讲第二章普通物理（一）.doc',
    #                r'E:\考证\岩土\新建文件夹\第3讲第二章普通物理（一）.pdf')
    # ll = [r'D:\下载\数理化讲义\第1讲高等数学考前宣讲.doc',
    #       r'D:\下载\数理化讲义\第2讲普通物理考前复习指导.doc']
    # Utils.mergewords(ll, r'E:\考证\岩土\新建文件夹\merged.doc')

    # 获得文档
    # file = docx.Document(ll[0])
    # file2 = docx.Document(ll[1])
    #
    # strText = []
    # # 按照段落读取文档内容
    # for para in file.paragraphs:
    #     strText.append(para.text)
    # for para2 in file2.paragraphs:
    #     strText.append(para2.text)
    # for s in strText:
    #     print(s)
